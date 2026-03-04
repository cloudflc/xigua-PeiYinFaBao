import re
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class ScriptConverter:
    def __init__(self):
        self.role_pattern = re.compile(r'^([^(]+)(?:\(([^)]+)\))?[:：]$')
        self.section_pattern = re.compile(r'^--------')
        self.excluded_keywords = ['简介', '事情', '景色', '温馨提示', '我们的任务是', '我们就进入第二步', '接着，我们掌握了', '通常有两种方式']
    
    def is_valid_role(self, role_name):
        for keyword in self.excluded_keywords:
            if keyword in role_name:
                return False
        
        if len(role_name) > 15:
            return False
        
        return True
    
    def parse_script(self, file_path):
        doc = Document(file_path)
        dialogues = []
        current_role = None
        current_position = None
        current_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            if self.section_pattern.match(text):
                if current_role and current_content:
                    content = '\n'.join(current_content).strip()
                    if content and '倩倩老师' not in current_role:
                        dialogues.append({
                            'role': current_role,
                            'position': current_position,
                            'content': content
                        })
                current_role = None
                current_position = None
                current_content = []
                continue
            
            match = self.role_pattern.match(text)
            if match:
                if current_role and current_content:
                    content = '\n'.join(current_content).strip()
                    if content and '倩倩老师' not in current_role:
                        dialogues.append({
                            'role': current_role,
                            'position': current_position,
                            'content': content
                        })
                
                role_name = match.group(1).strip()
                
                if not self.is_valid_role(role_name):
                    current_role = None
                    current_position = None
                    current_content = []
                    continue
                
                current_role = role_name
                current_position = match.group(2).strip() if match.group(2) else ""
                current_content = []
            elif current_role:
                current_content.append(text)
        
        if current_role and current_content:
            content = '\n'.join(current_content).strip()
            if content and current_role != '倩倩老师':
                dialogues.append({
                    'role': current_role,
                    'position': current_position,
                    'content': content
                })
        
        return dialogues
    
    def create_template(self, dialogues, output_file):
        doc = Document()
        
        title = doc.add_heading('配音发包模板', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        headers = ['序号', '角色', '台词', '备注']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
        
        for idx, dialogue in enumerate(dialogues, 1):
            row_cells = table.add_row().cells
            
            row_cells[0].text = str(idx)
            row_cells[1].text = dialogue['role']
            row_cells[2].text = dialogue['content']
            row_cells[3].text = ''
            
            for i in range(4):
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        doc.save(output_file)
        print(f"转换完成！输出文件：{output_file}")
        print(f"共转换 {len(dialogues)} 条对话")

def main():
    if len(sys.argv) < 2:
        print("使用方法: python script_converter.py <脚本文件.docx> [输出文件.docx]")
        print("示例: python script_converter.py 音乐播放器-1.docx")
        print("示例: python script_converter.py 音乐播放器-1.docx 输出模板.docx")
        return
    
    converter = ScriptConverter()
    
    input_file = sys.argv[1]
    
    if not os.path.exists(input_file):
        print(f"错误：文件不存在 - {input_file}")
        return
    
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}_配音发包模板.docx"
    
    print(f"正在读取脚本文件：{input_file}")
    dialogues = converter.parse_script(input_file)
    
    if not dialogues:
        print("警告：未找到任何对话内容")
        return
    
    print(f"正在生成配音发包模板：{output_file}")
    converter.create_template(dialogues, output_file)

if __name__ == "__main__":
    main()
