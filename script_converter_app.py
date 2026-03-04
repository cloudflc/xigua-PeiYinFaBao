import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from datetime import datetime


class ScriptConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("脚本转配音发包模板工具")
        self.root.geometry("900x700")
        
        self.script_file = None
        self.all_roles = []
        self.selected_roles = set()
        self.dialogues = []
        self.default_roles = ['小八', '孙小弟', '胖达', '妮可', '旁白']
        
        self.setup_styles()
        self.create_widgets()
    
    def setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        style.configure('Title.TLabel', font=('Microsoft YaHei UI', 18, 'bold'), foreground='#2c3e50')
        style.configure('Header.TLabel', font=('Microsoft YaHei UI', 11, 'bold'), foreground='#34495e')
        style.configure('Normal.TLabel', font=('Microsoft YaHei UI', 10), foreground='#555555')
        style.configure('Status.TLabel', font=('Microsoft YaHei UI', 9), foreground='#7f8c8d', relief=tk.SUNKEN)
        
        style.configure('Primary.TButton', font=('Microsoft YaHei UI', 10), padding=8)
        style.configure('Success.TButton', font=('Microsoft YaHei UI', 10, 'bold'), padding=8)
        
        style.configure('TFrame', background='#f8f9fa')
        style.configure('Card.TFrame', background='#ffffff', relief=tk.RAISED, borderwidth=1)
        
        self.root.configure(bg='#f8f9fa')
    
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(3, weight=1)
        
        title_label = ttk.Label(main_frame, text="📝 脚本转配音发包模板工具", style='Title.TLabel')
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 25))
        
        file_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15")
        file_card.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 20))
        
        ttk.Label(file_card, text="选择脚本文件：", style='Header.TLabel').grid(row=0, column=0, sticky=tk.W, pady=(0, 10))
        
        file_frame = ttk.Frame(file_card)
        file_frame.grid(row=1, column=0, sticky=(tk.W, tk.E))
        file_frame.columnconfigure(0, weight=1)
        
        self.file_entry = ttk.Entry(file_frame, width=60, font=('Microsoft YaHei UI', 10))
        self.file_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 10))
        
        ttk.Button(file_frame, text="📁 浏览", command=self.browse_file, style='Primary.TButton').grid(row=0, column=1)
        
        role_card = ttk.Frame(main_frame, style='Card.TFrame', padding="15")
        role_card.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 20))
        
        ttk.Label(role_card, text="选择需要配音的角色：", style='Header.TLabel').grid(row=0, column=0, sticky=tk.W, pady=(0, 10))
        
        role_frame = ttk.Frame(role_card)
        role_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        role_frame.columnconfigure(0, weight=1)
        
        self.role_listbox = tk.Listbox(role_frame, selectmode=tk.MULTIPLE, height=15, 
                                     font=('Microsoft YaHei UI', 10), bg='#ffffff',
                                     selectbackground='#3498db', selectforeground='#ffffff',
                                     relief=tk.FLAT, borderwidth=1)
        self.role_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        scrollbar = ttk.Scrollbar(role_frame, orient=tk.VERTICAL, command=self.role_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.role_listbox.config(yscrollcommand=scrollbar.set)
        
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=(0, 20))
        
        ttk.Button(button_frame, text="✓ 全选", command=self.select_all, style='Primary.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="✗ 全不选", command=self.deselect_all, style='Primary.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="🚀 生成模板", command=self.generate_template, style='Success.TButton').pack(side=tk.LEFT, padx=5)
        
        self.status_label = ttk.Label(main_frame, text="请选择脚本文件", style='Status.TLabel')
        self.status_label.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 0))
        
        self.generated_doc = None
        self.output_file = None
    
    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="选择脚本文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.script_file = file_path
            self.status_label.config(text=f"正在分析：{os.path.basename(file_path)}")
            self.root.update()
            self.analyze_file()
    
    def analyze_file(self):
        if not self.script_file:
            return
        
        try:
            doc = Document(self.script_file)
            role_pattern = re.compile(r'^([^(]+)(?:\(([^)]+)\))?[:：]$')
            self.all_roles = set()
            self.dialogues = []
            
            current_role = None
            current_position = None
            current_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                match = role_pattern.match(text)
                if match:
                    if current_role and current_content:
                        content = '\n'.join(current_content).strip()
                        content = self.remove_brackets_content(content)
                        if content:
                            self.dialogues.append({
                                'role': current_role,
                                'position': current_position,
                                'content': content
                            })
                    
                    current_role = match.group(1).strip()
                    current_position = match.group(2).strip() if match.group(2) else ""
                    current_content = []
                    self.all_roles.add(current_role)
                elif current_role:
                    current_content.append(text)
            
            if current_role and current_content:
                content = '\n'.join(current_content).strip()
                content = self.remove_brackets_content(content)
                if content:
                    self.dialogues.append({
                        'role': current_role,
                        'position': current_position,
                        'content': content
                    })
            
            self.all_roles = sorted(list(self.all_roles))
            
            self.role_listbox.delete(0, tk.END)
            for i, role in enumerate(self.all_roles):
                self.role_listbox.insert(tk.END, role)
                if self.should_default_select(role):
                    self.role_listbox.selection_set(i)
            
            self.status_label.config(text=f"✅ 分析完成！发现 {len(self.all_roles)} 个可能的角色，{len(self.dialogues)} 条对话")
            
        except Exception as e:
            messagebox.showerror("错误", f"分析文件时出错：{str(e)}")
            self.status_label.config(text="❌ 分析失败")
    
    def remove_brackets_content(self, text):
        pattern = r'[（\(][^）\)]*[）\)]'
        result = re.sub(pattern, '', text)
        return result.strip()
    
    def should_default_select(self, role):
        for default_role in self.default_roles:
            if default_role in role:
                return True
        return False
    
    def select_all(self):
        self.role_listbox.selection_set(0, tk.END)
    
    def deselect_all(self):
        self.role_listbox.selection_clear(0, tk.END)
    
    def generate_template(self):
        selected_indices = self.role_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("警告", "请至少选择一个角色！")
            return
        
        self.selected_roles = set()
        for idx in selected_indices:
            self.selected_roles.add(self.role_listbox.get(idx))
        
        filtered_dialogues = [d for d in self.dialogues if d['role'] in self.selected_roles]
        
        if not filtered_dialogues:
            messagebox.showwarning("警告", "选中的角色没有对话内容！")
            return
        
        default_name = os.path.splitext(os.path.basename(self.script_file))[0] + "_配音发包模板.docx"
        
        file_path = filedialog.asksaveasfilename(
            title="保存配音发包模板",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            self.create_template_document(filtered_dialogues, file_path)
            self.status_label.config(text=f"✅ 文件已保存：{os.path.basename(file_path)}")
            messagebox.showinfo("成功", f"🎉 模板生成完成！\n\n共 {len(filtered_dialogues)} 条对话\n已保存到：\n{file_path}")
        except Exception as e:
            messagebox.showerror("错误", f"保存文件时出错：{str(e)}")
    
    def create_template_document(self, dialogues, output_file):
        doc = Document()
        
        title = doc.add_paragraph(f"西瓜创客配音Python需求单-{os.path.splitext(os.path.basename(self.script_file))[0]}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(16)
            run.font.bold = True
        
        current_date = datetime.now().strftime("%m月%d日")
        doc.add_paragraph(f"时间 {current_date}")
        
        doc.add_paragraph("写在前面，开始录制前务必仔细阅读#")
        
        intro_text = """西瓜创客是一家在线少儿编程教育公司，用户对象为6-12岁小朋友，课程是以视频的形式在线上进行学习，本次配音是用于课程中的角色，需要保证声音塑造符合用户年龄层次，配音需要符合对应人物性别和性格特征。"""
        doc.add_paragraph(intro_text)
        
        doc.add_paragraph("版权及保密")
        doc.add_paragraph("凡牵涉甲方的所有文档、音视频文件、图片素材均需要保密，未经甲方授权，不可透露无关人员")
        
        doc.add_paragraph("角色说明")
        
        role_descriptions = {
            "孙小弟": "——正派\n男，心理年龄相当于人类的 5 / 6 岁。古腾堡学院的学员，活泼，可爱，淘气， 充满好奇心，喜欢捣乱，喜欢用调皮的招式对付敌人。去「飞船捣乱」就是他的主意。",
            "妮可": "——正派\n女性，古腾堡学院的学生，孙小弟的小伙伴之一。实际年龄相当于人类的 8 岁。聪明伶俐，性格外向，有着女孩子特有的温柔，善于思考和发现问题。非常爱美，怕脏。会一定魔法（但遇到事情时几乎不起作用）。",
            "小八": "孙小弟的小伙伴之一，一个拥有超级AI的小机器人，其人格对应的年龄为 5 岁左右。古腾堡学院的吉祥物，是非观非常非常明确，能够准确地看到事情中不对的一面，对其他人进行劝阻（虽然通常无效，并且尝尝被逼无奈做自己认为不对的事情）。有着超乎想象的知识储备量，是走动的百科全书。有人类的情感，喜欢夸奖，喜欢被摸头，讨厌一个人相处，讨厌被说自己没用，对妮可的魔法非常好奇。",
            "胖达": "——正派\n孙小弟的小伙伴之一，男，相当于年龄7岁。古腾堡学院的学生。贪吃，喜欢吃各种好吃的食物，听见食物就会流口水，看见食物就会不受控制地去吃。善良而单纯。讨厌运动，喜欢捯饬修理各种机器，但水平不高，修好一个毛病的同时，会造成一个新的毛病。"
        }
        
        for role in sorted(self.selected_roles):
            if role in role_descriptions:
                doc.add_paragraph(role)
                doc.add_paragraph(role_descriptions[role])
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        headers = ['人物', '台词', '备注']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
        
        for dialogue in dialogues:
            row_cells = table.add_row().cells
            row_cells[0].text = dialogue['role']
            row_cells[1].text = dialogue['content']
            row_cells[2].text = ''
            
            for i in range(3):
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        doc.save(output_file)


def main():
    root = tk.Tk()
    app = ScriptConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
