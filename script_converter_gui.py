import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os


class ScriptConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("脚本转配音发包模板工具")
        self.root.geometry("800x600")
        
        self.script_file = None
        self.all_roles = []
        self.selected_roles = set()
        self.dialogues = []
        
        self.create_widgets()
    
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(3, weight=1)
        
        ttk.Label(main_frame, text="脚本转配音发包模板工具", font=("Arial", 16, "bold")).grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        ttk.Label(main_frame, text="选择脚本文件：").grid(row=1, column=0, sticky=tk.W, pady=5)
        
        file_frame = ttk.Frame(main_frame)
        file_frame.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5)
        file_frame.columnconfigure(0, weight=1)
        
        self.file_entry = ttk.Entry(file_frame, width=50)
        self.file_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        ttk.Button(file_frame, text="浏览", command=self.browse_file).grid(row=0, column=1)
        ttk.Button(file_frame, text="分析", command=self.analyze_file).grid(row=0, column=2, padx=(5, 0))
        
        ttk.Label(main_frame, text="选择需要配音的角色：").grid(row=2, column=0, sticky=(tk.W, tk.N), pady=(15, 5))
        
        self.role_listbox = tk.Listbox(main_frame, selectmode=tk.MULTIPLE, height=15)
        self.role_listbox.grid(row=3, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(15, 5))
        
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.role_listbox.yview)
        scrollbar.grid(row=3, column=2, sticky=(tk.N, tk.S), pady=(15, 5))
        self.role_listbox.config(yscrollcommand=scrollbar.set)
        
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=2, pady=(20, 0))
        
        ttk.Button(button_frame, text="全选", command=self.select_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="全不选", command=self.deselect_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="生成模板", command=self.generate_template).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="保存", command=self.save_file).pack(side=tk.LEFT, padx=5)
        
        self.status_label = ttk.Label(main_frame, text="请选择脚本文件并点击分析", relief=tk.SUNKEN)
        self.status_label.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(20, 0))
        
        self.generated_doc = None
        self.output_file = None
    
    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="选择脚本文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.script_file = file_path
            self.status_label.config(text=f"已选择文件：{os.path.basename(file_path)}")
    
    def analyze_file(self):
        if not self.script_file:
            messagebox.showwarning("警告", "请先选择脚本文件！")
            return
        
        try:
            doc = Document(self.script_file)
            role_pattern = re.compile(r'^([^(]+)(?:\(([^)]+)\))?[:：]$')
            self.all_roles = set()
            self.dialogues = []
            
            current_role = None
            current_position = None
            current_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                match = role_pattern.match(text)
                if match:
                    if current_role and current_content:
                        content = '\n'.join(current_content).strip()
                        if content:
                            self.dialogues.append({
                                'role': current_role,
                                'position': current_position,
                                'content': content
                            })
                    
                    current_role = match.group(1).strip()
                    current_position = match.group(2).strip() if match.group(2) else ""
                    current_content = []
                    self.all_roles.add(current_role)
                elif current_role:
                    current_content.append(text)
            
            if current_role and current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    self.dialogues.append({
                        'role': current_role,
                        'position': current_position,
                        'content': content
                    })
            
            self.all_roles = sorted(list(self.all_roles))
            
            self.role_listbox.delete(0, tk.END)
            for role in self.all_roles:
                self.role_listbox.insert(tk.END, role)
            
            self.status_label.config(text=f"分析完成！发现 {len(self.all_roles)} 个可能的角色，{len(self.dialogues)} 条对话")
            
        except Exception as e:
            messagebox.showerror("错误", f"分析文件时出错：{str(e)}")
    
    def select_all(self):
        self.role_listbox.selection_set(0, tk.END)
    
    def deselect_all(self):
        self.role_listbox.selection_clear(0, tk.END)
    
    def generate_template(self):
        selected_indices = self.role_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("警告", "请至少选择一个角色！")
            return
        
        self.selected_roles = set()
        for idx in selected_indices:
            self.selected_roles.add(self.role_listbox.get(idx))
        
        filtered_dialogues = [d for d in self.dialogues if d['role'] in self.selected_roles]
        
        if not filtered_dialogues:
            messagebox.showwarning("警告", "选中的角色没有对话内容！")
            return
        
        self.generated_doc = Document()
        title = self.generated_doc.add_heading('配音发包模板', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table = self.generated_doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        headers = ['序号', '角色', '台词', '备注']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
        
        for idx, dialogue in enumerate(filtered_dialogues, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = dialogue['role']
            row_cells[2].text = dialogue['content']
            row_cells[3].text = ''
            
            for i in range(4):
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        self.status_label.config(text=f"模板生成完成！共 {len(filtered_dialogues)} 条对话，请点击保存")
        messagebox.showinfo("成功", f"模板生成完成！共 {len(filtered_dialogues)} 条对话")
    
    def save_file(self):
        if not self.generated_doc:
            messagebox.showwarning("警告", "请先生成模板！")
            return
        
        default_name = os.path.splitext(os.path.basename(self.script_file))[0] + "_配音发包模板.docx"
        
        file_path = filedialog.asksaveasfilename(
            title="保存配音发包模板",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        
        if file_path:
            try:
                self.generated_doc.save(file_path)
                self.output_file = file_path
                self.status_label.config(text=f"文件已保存：{os.path.basename(file_path)}")
                messagebox.showinfo("成功", f"文件已成功保存到：\n{file_path}")
            except Exception as e:
                messagebox.showerror("错误", f"保存文件时出错：{str(e)}")


def main():
    root = tk.Tk()
    app = ScriptConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
