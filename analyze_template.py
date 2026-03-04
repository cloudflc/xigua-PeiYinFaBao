from docx import Document

def analyze_template(file_path):
    doc = Document(file_path)
    
    print(f"=== {file_path} 完整内容分析 ===\n")
    
    print("【段落内容】")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"段落 {i}: {para.text.strip()}")
    
    print("\n【表格内容】")
    if doc.tables:
        table = doc.tables[0]
        print(f"表格行数: {len(table.rows)}")
        print(f"表格列数: {len(table.columns)}")
        
        print("\n表头:")
        for cell in table.rows[0].cells:
            print(f"  {cell.text}")
        
        print("\n前3行数据:")
        for i in range(1, min(4, len(table.rows))):
            row = table.rows[i]
            cells = [cell.text for cell in row.cells]
            print(f"  行{i}: {cells}")

analyze_template("配音发包模板.docx")
